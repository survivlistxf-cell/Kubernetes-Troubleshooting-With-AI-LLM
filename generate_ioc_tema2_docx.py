# -*- coding: utf-8 -*-
"""
Generates "IOC - Tema 2" style DOCX for the Kubexplain thesis project.
Uses python-docx for Document generation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class IOCDocxBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        # Default style for Normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(30, 30, 30)

    def add_title_page(self):
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Tema 2")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Proiectarea centrata pe sarcini\nsi prototipizarea")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        self.doc.add_paragraph() # Spacing

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Proiect: Kubexplain - Asistent AI pentru Diagnosticarea\nProblemelor in Clustere Kubernetes")
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(60, 60, 60)

        self.doc.add_paragraph() # Spacing

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Axinescu Valentin-Daniel")
        run.font.size = Pt(11)
        run.italic = True
        
        self.doc.add_page_break()

    def table_of_contents(self, items):
        p = self.doc.add_paragraph()
        run = p.add_run("Cuprins")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        for num, title, page in items:
            p = self.doc.add_paragraph()
            line = f"{num} {title}" if num else title
            
            run = p.add_run(line)
            p.paragraph_format.space_after = Pt(2)
            
            # Simple simulation of table of contents dots
            dots = "." * max(1, 100 - len(line) - len(page))
            p.add_run(f" {dots} {page}")

        self.doc.add_page_break()

    def chapter_title(self, title, level=1):
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 70, 130)
            p.paragraph_format.space_before = Pt(12)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(30, 30, 30)
            p.paragraph_format.space_before = Pt(10)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(50, 50, 50)
            p.paragraph_format.space_before = Pt(6)

        p.paragraph_format.space_after = Pt(6)

    def body_text(self, text):
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    def bullet(self, text, level=0):
        # style List Bullet (levels handle indentation automatically)
        p = self.doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        # Manually handling sub-bullets since nested bullets in python-docx can be tricky
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * level)

    def italic_text(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(60, 60, 60)
        p.paragraph_format.space_after = Pt(6)

    def bold_text(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.paragraph_format.space_after = Pt(4)

    def discussion_block(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(12)

    def separator(self):
        p = self.doc.add_paragraph()
        run = p.add_run("_" * 50)
        run.font.color.rgb = RGBColor(180, 180, 180)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, output_path):
        self.doc.save(output_path)


def build_docx():
    builder = IOCDocxBuilder()

    # =============================================
    # TITLE PAGE
    # =============================================
    builder.add_title_page()

    # =============================================
    # TABLE OF CONTENTS
    # =============================================
    toc_items = [
        ("1.", "Descrierea temei", "2"),
        ("2.", "Ce contine portofoliul", "2"),
        ("", "  Sectiunea 1 - Sarcini si cerinte", "3"),
        ("", "  Sectiunea 2 - Primul prototip si parcurgerea lui", "3"),
        ("3.", "Etapele rezolvarii temei", "3"),
        ("4.", "Sectiunea 1: Sarcini si cerinte - Kubexplain", "4"),
        ("", "  Introducere: cadrul general al sistemului", "4"),
        ("", "  Introducere: utilizatori potentiali", "5"),
        ("", "  Introducere: contexte de lucru", "5"),
        ("", "  Introducere: la ce va fi folosit sistemul", "6"),
        ("", "  Introducere: constrangerile sistemului", "6"),
        ("", "  Exemple de sarcini", "7"),
        ("", "  Lista de cerinte", "10"),
        ("5.", "Sectiunea 2: Prototip si parcurgere orientata pe sarcini", "12"),
        ("", "  Parcurgerea sarcinii 1", "12"),
        ("", "  Parcurgerea sarcinii 2", "14"),
        ("", "  Preocupari majore din parcurgerea sarcinilor", "15"),
        ("6.", "Bibliografie", "16"),
    ]
    builder.table_of_contents(toc_items)

    # =============================================
    # 1. DESCRIEREA TEMEI
    # =============================================
    builder.chapter_title("1. Descrierea temei")
    
    builder.body_text(
        "Prim pas in proiectarea iterativa a unui sistem: proiectarea centrata pe sarcini. "
        "Se incepe prin identificarea utilizatorilor si a sarcinilor pe care le au de indeplinit, "
        "precum si contextul desfasurarii actiunilor lor. Un sistem este proiectat asa cum trebuie "
        "numai daca ia in considerare oamenii, sarcinile si nevoile lor."
    )

    builder.body_text(
        "Tema propusa - inceperea unei proiectari iterative pentru sistemul Kubexplain, un asistent "
        "AI pentru diagnosticarea problemelor in clustere Kubernetes, folosind metoda proiectarii "
        "centrate pe sarcini si metode de prototipizare."
    )

    builder.body_text("Scop - capatarea de experienta in:")
    builder.bullet("Realizarea unei descrieri bune a sarcinilor utilizatorilor sistemului Kubexplain")
    builder.bullet("Folosirea descrierii sarcinilor pentru a decide cu privire la cerintele sistemului")
    builder.bullet("Realizarea unor prototipuri pe baza celor de mai sus")
    builder.bullet("Evaluarea prototipurilor prin parcurgerea orientata pe sarcini")

    # =============================================
    # 2. CE CONTINE PORTOFOLIUL
    # =============================================
    builder.chapter_title("2. Ce contine portofoliul")

    builder.bullet("O lista care descrie potentialii utilizatori si contextele in care ei lucreaza")
    builder.bullet("O lista a sarcinilor reprezentative pe care oamenii se asteapta sa fie facute")
    builder.bullet("O lista de prioritati pe cerintele sistemului")
    builder.bullet("Un prototip functional al interfetei")
    builder.bullet("O parcurgere a prototipului centrata pe sarcini")

    builder.chapter_title("Sectiunea 1 - Sarcini si cerinte (aprox. 10 pag)", level=2)
    builder.body_text(
        "1. Introducere. Descrierea, in termeni generali, a comportamentului sistemului: utilizatori "
        "potentiali, contextele lor de lucru, ce fel de sistem se asteapta sa foloseasca.\n"
        "2. Exemple de sarcini. 5-7 exemple concrete de sarcini cu proprietatile din specificatii.\n"
        "3. Lista de cerinte. Pe baza exemplelor de sarcini se extrag principalele cerinte de sistem "
        "si se introduc prioritati."
    )

    builder.chapter_title("Sectiunea 2 - Primul prototip si parcurgerea lui", level=2)
    builder.body_text(
        "1. Prototip - dezvoltarea prototipului care sa satisfaca cerintele principale.\n"
        "2. Discutii si parcurgere. Parcurgerea prototipului pe baza sarcinilor identificate, "
        "evaluand cat de bine interfata se potriveste cu viziunea utilizatorilor.\n"
        "3. Preocupari majore identificate in urma parcurgerii."
    )

    # =============================================
    # 3. ETAPELE REZOLVARII TEMEI
    # =============================================
    builder.chapter_title("3. Etapele rezolvarii temei")

    builder.chapter_title("Etapa 1. Generarea unei liste de utilizatori potentiali si a unei liste initiale de sarcini", level=3)
    builder.body_text(
        "S-au identificat utilizatorii potentiali ai sistemului Kubexplain prin analiza directa a "
        "fluxurilor de lucru DevOps si prin discutii cu colegi care administreaza clustere Kubernetes. "
        "S-au observat sarcinile curente pe care le realizeaza: diagnosticarea podurilor cu probleme, "
        "verificarea starii nodurilor, analiza logurilor si evenimentelor. Acestea au stat la baza "
        "listei initiale de sarcini."
    )

    builder.chapter_title("Etapa 2. Validarea sarcinilor", level=3)
    builder.body_text(
        "S-a verificat lista sarcinilor cu potentiali utilizatori (ingineri DevOps, administratori de "
        "sisteme, studenti cu experienta in Kubernetes). S-au adaugat sarcini suplimentare precum "
        "gestionarea istoricului conversatiilor si atasarea de fisiere de context."
    )

    builder.chapter_title("Etapa 3. Deciderea utilizatorilor cheie si a unei prime liste de cerinte", level=3)
    builder.body_text(
        "Din exemplele de sarcini si discutiile cu utilizatorii, s-au decis cerintele majore ale "
        "sistemului si s-au prioritizat in: a) cele care trebuie incluse neaparat; b) cele care "
        "trebuie incluse; c) cele care pot fi incluse; d) cele care se exclud."
    )

    builder.chapter_title("Etapa 4. Dezvoltarea prototipului", level=3)
    builder.body_text(
        "S-a dezvoltat un prototip functional al aplicatiei web Kubexplain, cu interfata chat, "
        "scanere pentru poduri si noduri, sistem de autentificare, istoric conversatii si "
        "integrare cu un LLM (Ollama/Llama 3.1) pentru generarea raspunsurilor."
    )

    builder.chapter_title("Etapa 5. Parcurgere orientata pe sarcini", level=3)
    builder.body_text(
        "S-a testat prototipul prin parcurgerea orientata pe sarcini, identificandu-se probleme "
        "potentiale de utilizare si oportunitati de imbunatatire."
    )

    # =============================================
    # 4. SECTIUNEA 1: SARCINI SI CERINTE - KUBEXPLAIN
    # =============================================
    builder.page_break()
    builder.chapter_title("4. Sectiunea 1: Sarcini si cerinte - Kubexplain")

    # --- Introducere: cadrul general ---
    builder.chapter_title("Introducere: cadrul general al sistemului", level=2)

    builder.body_text(
        "Kubexplain este o aplicatie web de tip asistent inteligent, conceputa pentru a ajuta "
        "inginerii DevOps si administratorii de sisteme sa diagnosticheze si sa rezolve probleme "
        "in clustere Kubernetes. Sistemul combina o interfata de chat conversational cu un model "
        "de limbaj (LLM) si mecanisme de colectare automata a datelor din cluster."
    )

    builder.bullet(
        "Contextul actual: Diagnosticarea problemelor in Kubernetes presupune utilizarea manuala "
        "a comenzilor kubectl (describe, logs, events, get -o json), interpretarea output-urilor "
        "complexe si corelarea informatiilor din surse multiple. Acest proces este consumator "
        "de timp si necesita experienta semnificativa."
    )
    builder.bullet(
        "Problema pe care o rezolva: Kubexplain automatizeaza colectarea datelor diagnostice, "
        "le analizeaza cu ajutorul unui LLM si ofera solutii contextuale bazate pe documentatia "
        "oficiala Kubernetes si pe experienta acumulata."
    )
    builder.bullet(
        "Arhitectura: Sistemul este compus din trei componente principale containerizate cu Docker: "
        "un frontend web (Node.js/Express), un backend API (Spring Boot/Java), si un server AI "
        "(Spring Boot/Java cu integrare Ollama). Baza de date PostgreSQL stocheaza utilizatorii, "
        "conversatiile si datele de context."
    )

    # --- Introducere: utilizatori potentiali ---
    builder.chapter_title("Introducere: utilizatori potentiali", level=2)

    builder.bullet(
        "Inginer DevOps cu experienta (utilizator principal): Persoana care administreaza "
        "clustere Kubernetes zilnic, cu experienta in kubectl si concepte Kubernetes. Foloseste "
        "Kubexplain pentru a accelera diagnosticarea si pentru a obtine sugestii pentru probleme "
        "mai putin intalnite. Nu are nevoie de instruire formala, fiind familiarizat cu terminologia."
    )
    builder.bullet(
        "Administrator de sisteme/SRE (Site Reliability Engineer): Responsabil cu mentinerea "
        "disponibilitatii serviciilor. Utilizeaza Kubexplain in situatii de urgenta (incidente) "
        "pentru a identifica rapid cauza unei defectiuni. Are experienta medie spre avansata "
        "in Kubernetes."
    )
    builder.bullet(
        "Dezvoltator software cu cunostinte de baza in Kubernetes: Persoana care deployeaza "
        "aplicatii pe cluster dar nu administreaza infrastructura. Foloseste Kubexplain cand "
        "aplicatia sa nu porneste sau nu functioneaza corect. Necesita explicatii mai detaliate "
        "si ghidare pas cu pas."
    )
    builder.bullet(
        "Student/practicant in domeniul cloud: Persoana aflata in procesul de invatare a "
        "Kubernetes. Foloseste Kubexplain ca instrument educational, punand intrebari si "
        "primind explicatii despre concepte si comenzi."
    )

    # --- Introducere: contexte de lucru ---
    builder.chapter_title("Introducere: contexte de lucru", level=2)

    builder.bullet(
        "Munca de rutina: Verificarea periodica a starii clusterului - scanarea podurilor si "
        "nodurilor, identificarea celor cu probleme (CrashLoopBackOff, Pending, ImagePullBackOff). "
        "Aceasta sarcina se realizeaza zilnic, de obicei la inceputul zilei de lucru."
    )
    builder.bullet(
        "Raspuns la incidente: Cand un serviciu cade sau un pod refuza sa porneasca, inginerul "
        "trebuie sa actioneze rapid. Kubexplain este folosit ca prim instrument de investigatie, "
        "oferind o analiza automata a logurilor si evenimentelor."
    )
    builder.bullet(
        "Invatare si explorare: Studentii si dezvoltatorii folosesc Kubexplain pentru a intelege "
        "mai bine conceptele Kubernetes, punand intrebari in limbaj natural si primind explicatii "
        "contextualizate."
    )
    builder.bullet(
        "Lucrul cu clustere remote: Sistemul este deployat pe un server cloud (OpenStack) si "
        "accesat prin browser, permitand diagnosticarea de la distanta fara acces direct la "
        "linia de comanda a clusterului."
    )

    # --- Introducere: la ce va fi folosit sistemul ---
    builder.chapter_title("Introducere: la ce va fi folosit sistemul", level=2)

    builder.body_text("Sistemul va gestiona urmatoarele activitati:")
    builder.bullet("Chat conversational cu AI pentru diagnosticarea problemelor Kubernetes")
    builder.bullet("Utilizatorul descrie problema in limbaj natural", level=1)
    builder.bullet("AI-ul analizeaza si ofera solutii bazate pe documentatia oficiala", level=1)
    builder.bullet("Suport pentru conversatii multi-tura cu pastrarea contextului", level=1)

    builder.bullet("Scanarea si inspectia resurselor din cluster")
    builder.bullet("Scanare poduri pe namespace cu vizualizare stare (Running, Pending, Failed)", level=1)
    builder.bullet("Scanare noduri cu detalii despre rol (Master/Worker) si resurse", level=1)
    builder.bullet("Vizualizare detaliata: Describe, JSON, Events, Logs pentru fiecare resursa", level=1)

    builder.bullet("Atasarea de context diagnostic la conversatie")
    builder.bullet("Adaugarea automata a datelor scanate (pod describe, logs) in conversatie", level=1)
    builder.bullet("Incarcarea de fisiere (YAML, logs, configuratii) ca atasamente", level=1)
    builder.bullet("Selectarea nivelului de detaliu (describe, JSON, events, logs)", level=1)

    builder.bullet("Gestionarea istoricului conversatiilor")
    builder.bullet("Salvarea automata a tuturor conversatiilor in baza de date", level=1)
    builder.bullet("Reluarea conversatiilor anterioare", level=1)
    builder.bullet("Editarea/stergerea/regenerarea titlului conversatiilor", level=1)

    builder.bullet("Autentificare si gestionarea conturilor de utilizator")
    builder.bullet("Inregistrare cu username, email si parola", level=1)
    builder.bullet("Login/Logout cu persistenta sesiunii", level=1)

    # --- Introducere: constrangerile sistemului ---
    builder.chapter_title("Introducere: constrangerile sistemului", level=2)

    builder.bullet(
        "Clusterul Kubernetes trebuie sa fie accesibil din containerul backend (prin montarea "
        "fisierului kubeconfig si a binarului kubectl)."
    )
    builder.bullet(
        "Serverul AI necesita un model LLM (Ollama cu Llama 3.1) rulat local sau pe un "
        "server dedicat cu resurse GPU/CPU suficiente."
    )
    builder.bullet(
        "Sistemul este deployat cu Docker Compose si necesita cel putin 4GB RAM total "
        "(2GB backend, 1GB frontend, 1GB+ pentru PostgreSQL si Ollama)."
    )
    builder.bullet(
        "Interfata web functioneaza in browser modern (Chrome, Firefox, Edge), fara "
        "instalare de software suplimentar pe client."
    )
    builder.bullet(
        "Baza de date PostgreSQL stocheaza toate datele persistent; nu se doreste inlocuirea, "
        "ci integrarea cu sistemele existente."
    )

    # =============================================
    # EXEMPLE DE SARCINI
    # =============================================
    builder.page_break()
    builder.chapter_title("Exemple de sarcini", level=2)
    builder.body_text(
        "Urmatoarele sarcini au fost colectate din observarea activitatilor inginerilor DevOps "
        "si a dezvoltatorilor care lucreaza cu Kubernetes, precum si din experienta proprie "
        "de administrare a clusterului licenta-cluster."
    )

    # --- SARCINA 1 ---
    builder.chapter_title("Sarcina 1: Diagnosticarea unui pod in CrashLoopBackOff", level=3)
    builder.body_text(
        "Andrei, inginer DevOps cu experienta de 3 ani, primeste o alerta ca serviciul "
        "backend-elearning din namespace-ul \"elearning\" nu mai raspunde. Deschide Kubexplain "
        "in browser si navigheaza la tab-ul \"Pods Scanner\". Introduce namespace-ul \"elearning\" "
        "si apasa \"Scan for Pods\". Vede ca podul backend-6f745b6fb5-jfsjv are statusul "
        "\"CrashLoopBackOff\" cu 15 restart-uri. Apasa pe pod pentru a vedea detaliile. "
        "In tab-ul \"Logs\" observa eroarea \"java.lang.OutOfMemoryError: Java heap space\". "
        "Selecteaza optiunea \"Add context to chat\" si se intoarce la chat. Scrie: \"De ce "
        "crashuieste podul meu de backend?\". AI-ul analizeaza logurile atasate si ii explica "
        "ca containerul depaseste limita de memorie si ii sugereaza sa creasca "
        "resources.limits.memory in deployment YAML."
    )
    builder.discussion_block(
        "Discutie: Aceasta este o sarcina frecventa si foarte importanta. Combina scanarea "
        "automata a clusterului cu analiza AI, reducand timpul de diagnosticare de la minute "
        "(manual cu kubectl) la secunde. Andrei este un utilizator tipic - experimentat dar care "
        "apreciaza automatizarea. Sarcina ilustreaza fluxul principal: Scan -> Inspect -> "
        "Add context -> Chat -> Solutie."
    )

    # --- SARCINA 2 ---
    builder.chapter_title("Sarcina 2: Investigarea unui pod blocat in Pending", level=3)
    builder.body_text(
        "Maria, administrator de sisteme, observa ca un nou deployment nu s-a finalizat. "
        "Podul postgres-backup-0 a ramas in starea \"Pending\" de 20 de minute. Deschide "
        "Kubexplain, scaneaza podurile din namespace-ul \"databases\" si identifica podul cu "
        "probleme. Apasa pe el, selecteaza tab-urile \"Describe\" si \"Events\", apoi adauga "
        "ambele ca context la chat. Intreaba: \"De ce nu porneste acest pod?\". AI-ul "
        "analizeaza evenimentele si identifica mesajul \"0/3 nodes are available: 3 Insufficient "
        "memory\". Ii explica ca niciun nod din cluster nu are suficienta memorie libera si "
        "ii sugereaza sa verifice consumul pe noduri sau sa adauge un nod nou."
    )
    builder.discussion_block(
        "Discutie: Sarcina demonstreaza o problema de scheduling frecvent intalnita. "
        "Maria utilizeaza multiple niveluri de detaliu (Describe + Events) pentru a oferi "
        "AI-ului context suficient. Sarcina este importanta si relativ frecventa, in special "
        "in clusterele cu resurse limitate."
    )

    # --- SARCINA 3 ---
    builder.chapter_title("Sarcina 3: Atasarea unui fisier YAML pentru analiza", level=3)
    builder.body_text(
        "Radu, dezvoltator junior, a scris un manifest YAML pentru deployment-ul aplicatiei "
        "sale dar primeste erori la \"kubectl apply\". Nu intelege ce este gresit. Deschide "
        "Kubexplain, se logheaza cu contul sau, apoi apasa butonul \"+\" de langa caseta de "
        "text si selecteaza \"Attach file\". Incarca fisierul deployment.yaml de pe calculatorul "
        "sau. Scrie in chat: \"Te rog analizeaza acest manifest YAML si spune-mi ce erori are\". "
        "AI-ul parseaza fisierul atasat si identifica doua probleme: un indent gresit la "
        "spec.containers si lipsa campului \"image\" pentru al doilea container. Radu corecteaza "
        "fisierul si deploy-ul functioneaza."
    )
    builder.discussion_block(
        "Discutie: Sarcina demonstreaza functionalitatea de upload de fisiere si analiza "
        "statica a configuratiilor. Este o sarcina frecventa pentru dezvoltatorii juniori. "
        "Ilustreaza si faptul ca utilizatorul poate sa nu aiba acces direct la cluster - "
        "poate lucra doar cu fisiere locale."
    )

    # --- SARCINA 4 ---
    builder.chapter_title("Sarcina 4: Verificarea starii nodurilor clusterului", level=3)
    builder.body_text(
        "Elena, SRE (Site Reliability Engineer), face verificarea de dimineata a clusterului. "
        "Deschide Kubexplain si navigheaza la \"Nodes Scanner\". Apasa \"Scan for Nodes\" si "
        "vede 3 noduri: un master si doi workeri. Observa ca nodul worker-2 are statusul "
        "\"NotReady\". Selecteaza \"Select Worker\" pentru a bifa doar nodurile worker, apoi "
        "alege nivelul de detaliu \"Describe\" + \"Events\" si apasa \"Add context for 2 "
        "selected nodes\". Contextul se adauga in chat. Intreaba: \"Ce se intampla cu worker-2? "
        "De ce e NotReady?\". AI-ul analizeaza output-ul de describe si identifica ca kubelet-ul "
        "nu mai comunica cu API server-ul de 15 minute. Sugereaza sa verifice daca nodul este "
        "accesibil prin SSH si daca serviciul kubelet functioneaza."
    )
    builder.discussion_block(
        "Discutie: Verificarea nodurilor este o sarcina de rutina efectuata zilnic. Scanner-ul "
        "de noduri ofera o vizualizare rapida a sanatatii intregului cluster. Selectarea in "
        "bulk a nodurilor si alegerea nivelului de detaliu demonstreaza eficienta interfetei "
        "pentru sarcini repetitive."
    )

    # --- SARCINA 5 ---
    builder.chapter_title("Sarcina 5: Reluarea unei conversatii anterioare", level=3)
    builder.body_text(
        "Andrei a investigat ieri o problema cu un serviciu si a primit o sugestie de la AI. "
        "Astazi vrea sa verifice daca solutia a functionat. Deschide Kubexplain, navigheaza la "
        "\"Chat History\" si gaseste conversatia cu titlul \"Backend OOM Investigation\". "
        "Apasa pe ea si conversatia anterioara se reincarca cu tot contextul. Scrie: \"Am marit "
        "limita la 2Gi dar tot crashuieste. Ce altceva pot face?\". AI-ul, avand contextul "
        "conversatiei anterioare, ii sugereaza sa verifice daca nu exista un memory leak in "
        "aplicatie si sa foloseasca un Java profiler."
    )
    builder.discussion_block(
        "Discutie: Aceasta sarcina demonstreaza persistenta conversatiilor si continuitatea "
        "contextului de-a lungul mai multor sesiuni. Este o functionalitate importanta "
        "care diferentiaza Kubexplain de un simplu chat AI - sistemul \"isi aminteste\" "
        "investigatiile anterioare. Sarcina este frecventa in cazul problemelor complexe "
        "care necesita mai multe iteratii de diagnosticare."
    )

    # --- SARCINA 6 ---
    builder.chapter_title("Sarcina 6: Adaugarea contextului scanat de poduri multiple", level=3)
    builder.body_text(
        "George, inginer DevOps, observa ca mai multe poduri din namespace-ul \"production\" "
        "au probleme simultan. Deschide Kubexplain, introduce namespace-ul \"production\" in "
        "Pods Scanner si apasa Scan. Vede 12 poduri, dintre care 4 au statusul diferit de "
        "\"Running\". Bifeaza \"Select All\" in bara de optiuni bulk, apoi deselecteaza "
        "manual podurile sanatoase. Alege nivelul de detaliu \"Describe\" + \"Logs\" si apasa "
        "\"Add context for 4 selected pods\". Intreaba in chat: \"Am mai multe poduri cu "
        "probleme simultan. Care este cauza comuna?\". AI-ul analizeaza cele 4 seturi de date "
        "si identifica un pattern comun: toate esueaza la conectarea la baza de date PostgreSQL. "
        "Sugereaza ca problema este la serviciul de baza de date, nu la podurile individuale."
    )
    builder.discussion_block(
        "Discutie: Aceasta sarcina complexa demonstreaza capacitatea sistemului de a corea "
        "informatii din surse multiple pentru a identifica o cauza radacina comuna. Selectarea "
        "in bulk si nivelurile de detaliu permit utilizatorului sa controleze cantitatea de "
        "context trimisa catre AI. Este o sarcina mai putin frecventa dar foarte importanta "
        "in situatii de outage."
    )

    # --- SARCINA 7 ---
    builder.chapter_title("Sarcina 7: Intrebare generala despre Kubernetes fara context de cluster", level=3)
    builder.body_text(
        "Ioana, studenta in anul 4, lucreaza la un proiect de semestru si nu intelege "
        "diferenta intre un Deployment si un StatefulSet in Kubernetes. Deschide Kubexplain "
        "si, fara sa scaneze vreun cluster, scrie direct in chat: \"Explica-mi diferenta "
        "dintre Deployment si StatefulSet si cand sa folosesc fiecare\". AI-ul ii ofera o "
        "explicatie detaliata, cu exemple de use-case-uri, folosind documentatia Kubernetes "
        "din baza de cunostinte (RAG). Ioana continua cu \"Poti sa imi dai un exemplu de "
        "StatefulSet YAML pentru un cluster Redis?\" si AI-ul genereaza un manifest complet."
    )
    builder.discussion_block(
        "Discutie: Aceasta sarcina demonstreaza ca Kubexplain nu este doar un instrument "
        "de diagnosticare, ci si un asistent educational. RAG-ul (Retrieval-Augmented "
        "Generation) permite raspunsuri bazate pe documentatia oficiala. Sarcina este "
        "frecventa pentru studentii si dezvoltatorii care invata Kubernetes."
    )

    # =============================================
    # LISTA DE CERINTE
    # =============================================
    builder.page_break()
    builder.chapter_title("Lista de cerinte", level=2)

    builder.chapter_title("Ce trebuie inclus neaparat (Must Have):", level=3)
    builder.bullet("Chat conversational cu AI (LLM) pentru diagnosticarea problemelor Kubernetes")
    builder.bullet("Scanarea podurilor dintr-un namespace specificat cu afisarea statusului")
    builder.bullet("Scanarea nodurilor clusterului cu afisarea rolului si statusului")
    builder.bullet("Vizualizare detaliata a resurselor: Describe, JSON, Events, Logs")
    builder.bullet("Adaugarea datelor scanate ca context in conversatia de chat")
    builder.bullet("Persistenta conversatiilor in baza de date PostgreSQL")
    builder.bullet("Sistem de autentificare (Register/Login)")
    builder.bullet("Protocol propriu de comunicare (kdiag/1.0) intre backend si serverul AI")
    builder.bullet("Integrare cu un LLM local (Ollama cu Llama 3.1)")

    builder.body_text("")
    builder.chapter_title("Ce trebuie inclus (Should Have):", level=3)
    builder.bullet("Istoric complet al conversatiilor cu posibilitatea de reluare")
    builder.bullet("Upload de fisiere (YAML, logs, configuratii) ca atasamente la conversatie")
    builder.bullet("RAG (Retrieval-Augmented Generation) cu documentatie Kubernetes scraping")
    builder.bullet("Dynamic RAG - cautare automata pe kubernetes.io cand LLM-ul nu gaseste raspuns")
    builder.bullet("Generare automata de titlu pentru conversatii (folosind LLM)")
    builder.bullet("Selectare in bulk a podurilor/nodurilor cu niveluri de detaliu configurabile")
    builder.bullet("Fallback heuristic cand serverul AI nu este disponibil")
    builder.bullet("Sistem de feedback (rating) pentru raspunsurile AI")

    builder.body_text("")
    builder.chapter_title("Ce ar putea fi inclus (Could Have):", level=3)
    builder.bullet("Dashboard cu metrici de performanta a clusterului in timp real")
    builder.bullet("Notificari automate cand un pod intra in stare de eroare")
    builder.bullet("Suport pentru mai multe modele LLM (selectie din Settings)")
    builder.bullet("Export conversatii ca PDF sau Markdown")
    builder.bullet("Integrare cu sisteme de alertare existente (Prometheus, Grafana)")
    builder.bullet("Dark/Light mode toggle complet functional")
    builder.bullet("Autocomplete pentru namespace-uri si nume de resurse")

    builder.body_text("")
    builder.chapter_title("Exclus:", level=3)
    builder.bullet("Executarea automata de comenzi kubectl de remediere fara confirmare umana")
    builder.bullet("Administrarea directa a clusterului (creare/stergere resurse) prin interfata")
    builder.bullet("Suport multi-cluster simultan in aceeasi sesiune")
    builder.bullet("Analiza de securitate sau audit RBAC")

    builder.italic_text(
        "Justificare: Executarea automata de comenzi de remediere a fost exclusa deliberat "
        "din motive de siguranta - un sistem AI nu trebuie sa faca modificari neautorizate "
        "intr-un cluster de productie. Administrarea directa depaseste scopul unui instrument "
        "de diagnosticare. Multi-cluster-ul si auditul de securitate ar creste complexitatea "
        "semnificativ fara a aduce valoare pentru MVP."
    )

    # =============================================
    # 5. SECTIUNEA 2: PROTOTIP SI PARCURGERE
    # =============================================
    builder.page_break()
    builder.chapter_title("5. Sectiunea 2: Prototip si parcurgere orientata pe sarcini")

    builder.body_text(
        "Prototipul Kubexplain este o aplicatie web functionala, implementata cu urmatorul "
        "stack tehnologic:"
    )
    builder.bullet("Frontend: HTML5 + CSS3 + JavaScript vanilla, servit prin Express.js (Node.js)")
    builder.bullet("Backend: Spring Boot (Java 17) cu API REST")
    builder.bullet("Server AI: Spring Boot (Java 17) cu integrare Ollama (Llama 3.1)")
    builder.bullet("Baza de date: PostgreSQL 16")
    builder.bullet("Containerizare: Docker Compose cu 3 servicii + baza de date")
    builder.bullet("Protocol: kdiag/1.0 - protocol propriu pentru comunicarea backend-AI server")

    builder.body_text(
        "Interfata principala este de tip single-page application cu sidebar de navigare "
        "si cinci tab-uri: Home (Chat), Chat History, Pods Scanner, Nodes Scanner si Settings."
    )

    builder.separator()

    # --- Parcurgerea Sarcinii 1 ---
    builder.chapter_title("Parcurgerea sarcinii 1: Diagnosticarea unui pod in CrashLoopBackOff", level=2)
    
    builder.chapter_title("Subsarcina 1: Autentificarea in sistem", level=3)
    builder.body_text(
        "a) Utilizatorul apasa butonul \"Login\" din sidebar-ul stang. Se deschide un modal "
        "cu formularul de autentificare.\n"
        "b) Introduce email-ul si parola. Apasa \"Login\".\n"
        "c) Daca credentialele sunt corecte, modalul se inchide, butonul \"Login\" se "
        "transforma in \"Logout\" si apare username-ul utilizatorului."
    )

    builder.chapter_title("Subsarcina 2: Scanarea podurilor", level=3)
    builder.body_text(
        "a) Se navigheaza la tab-ul \"Pods Scanner\" din sidebar.\n"
        "b) In campul \"Namespace\" se introduce \"elearning\".\n"
        "c) Se apasa butonul \"Scan for Pods\".\n"
        "d) Sistemul executa \"kubectl get pods -n elearning -o json\" pe server "
        "si afiseaza lista podurilor cu informatii: nume, status, restart-uri, varsta.\n"
        "e) Podurile cu probleme sunt evidentiate vizual (culori diferite pentru statusuri: "
        "verde pentru Running, rosu pentru CrashLoopBackOff, galben pentru Pending)."
    )

    builder.chapter_title("Subsarcina 3: Inspectia detaliata a podului", level=3)
    builder.body_text(
        "a) Se apasa pe podul backend-6f745b6fb5-jfsjv cu status CrashLoopBackOff.\n"
        "b) Se deschide un modal cu detalii complete, organizate in tab-uri: "
        "Describe, JSON, Events, Logs.\n"
        "c) In tab-ul Logs, se observa eroarea OOM.\n"
        "d) Se apasa \"Add context to chat\" - datele selectate se adauga automat "
        "ca artefact in conversatia curenta."
    )

    builder.chapter_title("Subsarcina 4: Conversatia cu AI-ul", level=3)
    builder.body_text(
        "a) Se revine la tab-ul Home (chat). In zona de preview a atasamentelor "
        "se vede contextul adaugat.\n"
        "b) Se scrie intrebarea: \"De ce crashuieste podul meu de backend?\".\n"
        "c) Se apasa Send (sau Enter). Mesajul + atasamentele sunt trimise la backend "
        "(port 8080), care le forwardeaza catre AI Server (port 8090) folosind protocolul "
        "kdiag/1.0.\n"
        "d) AI Server construieste un prompt de sistem cu documentatia Kubernetes relevanta "
        "(RAG), adauga mesajul utilizatorului cu artefactele, si trimite totul catre "
        "Ollama/Llama 3.1.\n"
        "e) Raspunsul AI-ului apare in zona de mesaje, formatat cu Markdown: paragrafe "
        "boldate pentru cauza, sugestii numerotate pentru solutii."
    )

    builder.separator()

    # --- Parcurgerea Sarcinii 2 ---
    builder.chapter_title("Parcurgerea sarcinii 2: Reluarea unei conversatii si adaugarea de fisier", level=2)

    builder.chapter_title("Subsarcina 1: Accesarea istoricului", level=3)
    builder.body_text(
        "a) Se navigheaza la tab-ul \"Chat History\".\n"
        "b) Se afiseaza lista conversatiilor anterioare cu titlu, data si preview.\n"
        "c) Se apasa pe conversatia dorita. Mesajele anterioare se reincarca in zona de chat."
    )

    builder.chapter_title("Subsarcina 2: Atasarea unui fisier", level=3)
    builder.body_text(
        "a) Se apasa butonul \"+\" de langa campul de text.\n"
        "b) Se selecteaza \"Attach file\" din dropdown.\n"
        "c) Se deschide selectorul de fisiere. Se alege deployment.yaml.\n"
        "d) Fisierul apare in zona de preview a atasamentelor cu numele, tipul MIME "
        "si dimensiunea.\n"
        "e) Se scrie intrebarea si se trimite. Fisierul este serializat si trimis "
        "ca artefact in payload-ul kdiag/1.0."
    )

    builder.chapter_title("Subsarcina 3: Gestionarea conversatiei", level=3)
    builder.body_text(
        "a) Titlul conversatiei este generat automat de LLM dupa primul mesaj.\n"
        "b) Utilizatorul poate edita titlul apasand butonul de editare.\n"
        "c) Se poate sterge conversatia din istoric cu confirmare.\n"
        "d) Se poate regenera titlul automat."
    )

    builder.separator()

    # =============================================
    # PREOCUPARI MAJORE
    # =============================================
    builder.chapter_title("Preocupari majore din parcurgerea sarcinilor de indeplinit", level=2)

    builder.bold_text("Fluxul Scan -> Context -> Chat necesita mai putini pasi")
    builder.bullet(
        "In prezent, utilizatorul trebuie sa: (1) navigheze la Pods Scanner, (2) scaneze, "
        "(3) apese pe pod, (4) apese \"Add context\", (5) revina la chat, (6) scrie "
        "intrebarea. Fluxul ar putea fi simplificat prin adaugarea unui buton \"Quick Debug\" "
        "direct pe fiecare pod din lista."
    )
    builder.body_text("") # spacing
    
    builder.bold_text("Cantitatea de context poate depasi limita LLM-ului")
    builder.bullet(
        "Cand se adauga context pentru 4+ poduri cu Describe + Logs, payload-ul poate "
        "depasi context window-ul LLM-ului. Sistemul trunchiaza automat la 10.000 "
        "caractere per artefact, dar utilizatorul nu este avertizat. Se sugereaza "
        "adaugarea unui indicator vizual al dimensiunii contextului."
    )
    builder.body_text("") # spacing

    builder.bold_text("Latenta raspunsurilor AI")
    builder.bullet(
        "Cu Llama 3.1 rulat local pe CPU, raspunsurile pot dura 30-60 secunde. "
        "In prototipul actual, utilizatorul vede doar un spinner fara progres. "
        "Se sugereaza implementarea streaming-ului de raspunsuri (token by token) "
        "pentru o experienta mai responsiva."
    )
    builder.body_text("") # spacing

    builder.bold_text("Feedback-ul si calitatea raspunsurilor")
    builder.bullet(
        "Sistemul de feedback (rating) permite colectarea aprecierilor, dar nu "
        "exista inca un mecanism de a folosi aceste date pentru imbunatatirea "
        "raspunsurilor. Se sugereaza implementarea unui sistem de RAG care sa "
        "salveze rezolvarile validate si sa le refoloseasca in cazuri similare."
    )
    builder.body_text("") # spacing

    builder.bold_text("Se sugereaza redesign minor al interfetei pentru:")
    builder.bullet("Integrarea mai fluida a scanner-elor in fluxul de chat")
    builder.bullet("Afisarea starii clusterului intr-un dashboard compact pe pagina principala")
    builder.bullet("Indicatori vizuali pentru loading/streaming al raspunsurilor AI")

    # =============================================
    # BIBLIOGRAFIE
    # =============================================
    builder.page_break()
    builder.chapter_title("6. Bibliografie")

    refs = [
        "1. Kubernetes Documentation. [Online] https://kubernetes.io/docs/",
        "2. Ollama - Run Large Language Models Locally. [Online] https://ollama.com/",
        "3. Spring Boot Reference Documentation. [Online] https://spring.io/projects/spring-boot",
        "4. Express.js Documentation. [Online] https://expressjs.com/",
        "5. PostgreSQL Documentation. [Online] https://www.postgresql.org/docs/",
        "6. Docker Compose Documentation. [Online] https://docs.docker.com/compose/",
        "7. Norman, Donald A. \"The Design of Everyday Things.\" Basic Books, 2013.",
        "8. Greenberg, Saul. CPSC 481 Human Computer Interaction I: Principles and Design. University of Calgary. [Online] http://pages.cpsc.ucalgary.ca/~saul/481/",
        "9. Lewis, C. & Rieman, J. Task-centered User Interface Design. 1993.",
        "10. Meta AI. \"Llama 3.1 - Open Foundation and Fine-Tuned Chat Models.\" 2024.",
    ]

    for ref in refs:
        builder.body_text(ref)

    # Save DOCX
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "IOC_Tema2_Kubexplain_Axinescu.docx")
    builder.save(output_path)
    print(f"DOCX generat cu succes: {output_path}")
    return output_path


if __name__ == "__main__":
    build_docx()
